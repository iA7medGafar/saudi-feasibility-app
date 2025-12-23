import streamlit as st
from google import genai
import pandas as pd
import time
import requests
from streamlit_lottie import st_lottie
import json
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from io import BytesIO
import random

# ==============================================================================
# 1. إعدادات النظام والتصميم (System Configuration)
# ==============================================================================
st.set_page_config(
    page_title="Jadwa AI | منصة جدوى",
    page_icon="💎",
    layout="wide",
    initial_sidebar_state="expanded"
)

# نظام الألوان والثيم (Design System)
# تم استخدام CSS Variables لسهولة التعديل وضمان التناسق
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@300;500;700;900&display=swap');

    :root {
        --primary-color: #4facfe;
        --secondary-color: #00f2fe;
        --bg-color: #0f172a;
        --card-bg: rgba(30, 41, 59, 0.7);
        --text-color: #f8fafc;
        --accent-green: #10b981;
        --accent-red: #ef4444;
    }

    /* تطبيق الخط العربي وتوحيد الاتجاه */
    * {
        font-family: 'Tajawal', sans-serif !important;
    }
    .stApp {
        background-color: var(--bg-color);
        background-image: radial-gradient(at 0% 0%, hsla(253,16%,7%,1) 0, transparent 50%), 
                          radial-gradient(at 50% 0%, hsla(225,39%,30%,1) 0, transparent 50%), 
                          radial-gradient(at 100% 0%, hsla(339,49%,30%,1) 0, transparent 50%);
        color: var(--text-color);
    }

    /* نصوص العناوين */
    h1, h2, h3, h4, p, span, div, label {
        direction: rtl;
        text-align: right;
        color: var(--text-color) !important;
    }

    /* البطاقات الزجاجية المحسنة (Frosted Glass) */
    .glass-container {
        background: var(--card-bg);
        backdrop-filter: blur(12px);
        -webkit-backdrop-filter: blur(12px);
        border: 1px solid rgba(255, 255, 255, 0.08);
        border-radius: 24px;
        padding: 30px;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.2);
        margin-bottom: 25px;
        transition: transform 0.3s ease;
    }
    .glass-container:hover {
        border-color: rgba(255, 255, 255, 0.2);
    }

    /* تخصيص حقول الإدخال */
    .stTextInput input, .stSelectbox div[data-baseweb="select"], .stNumberInput input, .stTextArea textarea {
        background-color: rgba(15, 23, 42, 0.6) !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
        color: white !important;
        border-radius: 12px !important;
    }

    /* الزر الرئيسي (Glow Effect) */
    .stButton > button {
        background: linear-gradient(135deg, var(--primary-color), var(--secondary-color));
        color: #0f172a !important;
        font-weight: 900 !important;
        border: none;
        border-radius: 50px;
        padding: 15px 40px;
        font-size: 1.2rem;
        width: 100%;
        box-shadow: 0 0 20px rgba(79, 172, 254, 0.4);
        transition: all 0.3s ease;
    }
    .stButton > button:hover {
        transform: scale(1.02);
        box-shadow: 0 0 30px rgba(79, 172, 254, 0.6);
    }

    /* بطاقات SWOT الاحترافية */
    .swot-card {
        padding: 20px;
        border-radius: 16px;
        height: 100%;
        border: 1px solid rgba(255,255,255,0.1);
        text-align: right;
    }
    .swot-s { background: linear-gradient(135deg, rgba(16, 185, 129, 0.2), rgba(16, 185, 129, 0.05)); border-left: 5px solid #10b981; }
    .swot-w { background: linear-gradient(135deg, rgba(239, 68, 68, 0.2), rgba(239, 68, 68, 0.05)); border-left: 5px solid #ef4444; }
    .swot-o { background: linear-gradient(135deg, rgba(59, 130, 246, 0.2), rgba(59, 130, 246, 0.05)); border-left: 5px solid #3b82f6; }
    .swot-t { background: linear-gradient(135deg, rgba(245, 158, 11, 0.2), rgba(245, 158, 11, 0.05)); border-left: 5px solid #f59e0b; }

    /* إخفاء عناصر Streamlit الافتراضية */
    #MainMenu {visibility: hidden;}
    header {visibility: hidden;}
    footer {visibility: hidden;}

</style>
""", unsafe_allow_html=True)

# ==============================================================================
# 2. وظائف المساعدة (Utility Functions) - العقل الهندسي
# ==============================================================================

@st.cache_data # استخدام الكاش لتسريع التحميل
def load_lottie(url: str):
    """تحميل ملفات الأنيميشن مع معالجة الأخطاء"""
    try:
        r = requests.get(url, timeout=3)
        if r.status_code != 200: return None
        return r.json()
    except: return None

def get_gemini_client():
    """جلب مفتاح API بأمان"""
    try:
        return st.secrets["GEMINI_KEY"]
    except:
        return None

def generate_with_retry(client, model, prompt, retries=3):
    """
    الدالة السحرية للتعامل مع أخطاء السيرفر (Exponential Backoff).
    إذا فشل الطلب، تنتظر وتجرب مرة أخرى تلقائياً.
    """
    for i in range(retries):
        try:
            response = client.models.generate_content(
                model=model,
                contents=prompt
            )
            return response.text
        except Exception as e:
            error_msg = str(e)
            if "429" in error_msg or "Resource has been exhausted" in error_msg:
                wait_time = (2 ** i) + random.uniform(0, 1) # انتظار ذكي: 1s, 2s, 4s...
                time.sleep(wait_time)
                continue # إعادة المحاولة
            else:
                raise e # إذا كان خطأ آخر، أوقفه
    raise Exception("عذراً، الخوادم مشغولة جداً حالياً. يرجى المحاولة لاحقاً.")

def create_professional_doc(data):
    """توليد ملف Word احترافي يدعم اللغة العربية بشكل صحيح"""
    doc = Document()
    
    # تنسيق الخط الافتراضي
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # العنوان
    title = doc.add_heading(f"دراسة جدوى: {data.get('project_name', 'مشروع')}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # دالة مساعدة لإضافة فقرات عربية
    def add_arabic_paragraph(text, style='Normal'):
        p = doc.add_paragraph(text, style=style)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # تفعيل RTL (من اليمين لليسار) للفقرة
        p.paragraph_format.bidi = True

    add_arabic_paragraph(data.get('summary', ''))
    
    doc.add_heading('تحليل SWOT', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    swot = data.get('swot', {})
    add_arabic_paragraph(f"نقاط القوة: {swot.get('s')}", 'List Bullet')
    add_arabic_paragraph(f"نقاط الضعف: {swot.get('w')}", 'List Bullet')
    add_arabic_paragraph(f"الفرص: {swot.get('o')}", 'List Bullet')
    add_arabic_paragraph(f"التهديدات: {swot.get('t')}", 'List Bullet')

    doc.add_heading('خطة العمل', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_arabic_paragraph(data.get('plan', ''))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==============================================================================
# 3. واجهة المستخدم (User Interface)
# ==============================================================================

# تحميل الأصول (Assets)
lottie_main = load_lottie("https://lottie.host/5b632675-5735-4d37-8898-33306db02b70/3k8l9z6j7a.json")
lottie_loading = load_lottie("https://lottie.host/98c2e061-0027-4c3e-b762-12711827453d/k1Y5g1o5mF.json")

# الهيدر (Header)
c1, c2 = st.columns([0.7, 0.3])
with c1:
    st.markdown("<h1 style='font-size: 3.5rem; margin-bottom: 0;'>💎 منصة جدوى</h1>", unsafe_allow_html=True)
    st.markdown("<p style='font-size: 1.2rem; opacity: 0.8;'>الجيل القادم من دراسات الجدوى المدعومة بالذكاء الاصطناعي.</p>", unsafe_allow_html=True)
with c2:
    if lottie_main: st_lottie(lottie_main, height=200, key="main_anim")

# نموذج الإدخال (Input Section)
st.markdown("<div class='glass-container'>", unsafe_allow_html=True)
col_input1, col_input2 = st.columns(2)

with col_input1:
    project_name = st.text_input("💡 اسم المشروع", placeholder="مثال: محمصة قهوة مختصة")
    capital = st.number_input("💰 رأس المال (SAR)", value=100000, step=10000, format="%d")

with col_input2:
    city = st.selectbox("📍 المدينة المستهدفة", ["الرياض", "جدة", "الدمام", "مكة المكرمة", "المدينة المنورة", "الخبر", "أخرى"])
    details = st.text_area("📝 تفاصيل إضافية (اختياري)", placeholder="ما الذي يميز مشروعك؟ من هم عملاؤك؟")

st.markdown("<br>", unsafe_allow_html=True)
analyze_btn = st.button("🚀 بدء التحليل الذكي")
st.markdown("</div>", unsafe_allow_html=True)

# ==============================================================================
# 4. محرك المنطق والذكاء الاصطناعي (Core Logic)
# ==============================================================================

if analyze_btn:
    api_key = get_gemini_client()
    
    if not api_key:
        st.error("⚠️ مفتاح API غير موجود. يرجى إضافته في إعدادات Secrets.")
    elif not project_name:
        st.warning("⚠️ يرجى إدخال اسم المشروع أولاً.")
    else:
        # واجهة التحميل التفاعلية (Status Container)
        status_container = st.status("جاري الاتصال بالمستشار الذكي...", expanded=True)
        
        try:
            client = genai.Client(api_key=api_key)
            
            # 1. مرحلة تحليل السوق
            status_container.write("🔍 جاري تحليل السوق والمنافسين...")
            if lottie_loading: 
                with st.columns([1,2,1])[1]: st_lottie(lottie_loading, height=120, key="proc")
            
            # بناء الموجه (Prompt) الدقيق جداً لضمان هيكلية JSON
            prompt = f"""
            أنت خبير اقتصادي ومستشار استراتيجي. قم بإعداد دراسة جدوى لمشروع "{project_name}" في مدينة "{city}" برأس مال {capital} ريال.
            التفاصيل الإضافية: {details}

            المهمة: أخرج النتيجة بصيغة JSON فقط وحصرياً. لا تكتب أي مقدمات ولا تستخدم Markdown (```json).
            يجب أن يكون الـ JSON بهذا الهيكل تماماً:
            {{
                "project_name": "{project_name}",
                "summary": "ملخص تنفيذي احترافي وجذاب لا يقل عن 5 أسطر",
                "swot": {{
                    "s": "نقطة قوة رئيسية",
                    "w": "نقطة ضعف رئيسية",
                    "o": "فرصة نمو في السوق",
                    "t": "تهديد محتمل"
                }},
                "financials": {{
                    "years": ["2025", "2026", "2027"],
                    "revenue": [150000, 250000, 400000],
                    "profit": [20000, 60000, 120000],
                    "notes": "ملاحظة مالية قصيرة"
                }},
                "plan": "خطة عمل تشغيلية وتسويقية مفصلة في فقرة واحدة متماسكة"
            }}
            """

            # 2. مرحلة المعالجة مع "إعادة المحاولة" (The Robust Call)
            status_container.write("🧠 جاري التفكير ومعالجة البيانات المالية...")
            
            # استخدام الموديل المستقر
            raw_response = generate_with_retry(client, 'gemini-1.5-flash', prompt)
            
            # 3. مرحلة التنظيف (Parsing)
            status_container.write("📊 تنسيق التقرير النهائي...")
            
            # تنظيف النص من أي شوائب (Markdown cleaning)
            clean_json = raw_response.replace("```json", "").replace("```", "").strip()
            data = json.loads(clean_json) # تحويل النص إلى كائن بايثون

            status_container.update(label="✅ اكتمل التحليل بنجاح!", state="complete", expanded=False)

            # =========================================================
            # عرض النتائج (Dashboard Layout)
            # =========================================================
            
            st.markdown("---")
            
            # التبويبات (Tabs)
            tab_overview, tab_swot, tab_finance, tab_plan = st.tabs([
                "📄 نظرة عامة", "⚖️ تحليل SWOT", "📈 المؤشرات المالية", "⚙️ خطة العمل"
            ])

            with tab_overview:
                st.markdown(f"<div class='glass-container'><h3>الملخص التنفيذي</h3><p>{data['summary']}</p></div>", unsafe_allow_html=True)

            with tab_swot:
                swot = data.get('swot', {})
                col_s, col_w, col_o, col_t = st.columns(4)
                with col_s: st.markdown(f"<div class='swot-card swot-s'><h4>💪 القوة</h4><p>{swot.get('s')}</p></div>", unsafe_allow_html=True)
                with col_w: st.markdown(f"<div class='swot-card swot-w'><h4>⚠️ الضعف</h4><p>{swot.get('w')}</p></div>", unsafe_allow_html=True)
                with col_o: st.markdown(f"<div class='swot-card swot-o'><h4>🌟 الفرص</h4><p>{swot.get('o')}</p></div>", unsafe_allow_html=True)
                with col_t: st.markdown(f"<div class='swot-card swot-t'><h4>🛡️ التهديدات</h4><p>{swot.get('t')}</p></div>", unsafe_allow_html=True)

            with tab_finance:
                fin = data.get('financials', {})
                col_chart, col_metrics = st.columns([2, 1])
                
                with col_chart:
                    st.markdown("<div class='glass-container'>", unsafe_allow_html=True)
                    df = pd.DataFrame({
                        "السنة": fin.get("years", []),
                        "الإيرادات": fin.get("revenue", []),
                        "الأرباح": fin.get("profit", [])
                    })
                    st.bar_chart(df.set_index("السنة"), color=["#4facfe", "#00f2fe"])
                    st.markdown("</div>", unsafe_allow_html=True)
                
                with col_metrics:
                    total_profit = sum(fin.get("profit", []))
                    roi = round((total_profit / capital) * 100, 1) if capital > 0 else 0
                    
                    st.markdown(f"""
                    <div class='glass-container' style='text-align: center;'>
                        <h4 style='margin:0'>إجمالي الربح (3 سنوات)</h4>
                        <h2 style='color: #10b981 !important;'>{total_profit:,} SAR</h2>
                        <hr style='border-color: rgba(255,255,255,0.1);'>
                        <h4 style='margin:0'>العائد على الاستثمار</h4>
                        <h2 style='color: #4facfe !important;'>{roi}%</h2>
                    </div>
                    """, unsafe_allow_html=True)

            with tab_plan:
                st.markdown(f"<div class='glass-container'><h3>خطة العمل المقترحة</h3><p>{data['plan']}</p></div>", unsafe_allow_html=True)

            # زر التحميل
            word_file = create_professional_doc(data)
            st.markdown("<br>", unsafe_allow_html=True)
            st.download_button(
                label="📥 تحميل الدراسة كاملة (Word Document)",
                data=word_file,
                file_name=f"Jadwa_{project_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
            st.balloons()

        except Exception as e:
            # معالجة الأخطاء بشكل لطيف للمستخدم
            error_message = str(e)
            if "JSON" in error_message:
                st.error("⚠️ حدث خطأ في معالجة البيانات القادمة من الذكاء الاصطناعي. يرجى المحاولة مرة أخرى.")
            elif "429" in error_message:
                st.warning("🚦 النظام تحت ضغط عالٍ حالياً. تم تفعيل نظام الانتظار الذكي، لكن المحاولة فشلت. جرب بعد دقيقة.")
            else:
                st.error(f"❌ حدث خطأ غير متوقع: {error_message}")
